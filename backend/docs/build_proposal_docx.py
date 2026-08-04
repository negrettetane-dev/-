from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "智途云枢项目策划书-新版.md"
OUTPUT = ROOT / "docs" / "智途云枢项目策划书-新版.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(91, 91, 91)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "F4F6F9"
BORDER = "D9E2F3"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
    else:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table):
    table.style = "Table Grid"
    cols = len(table.columns)
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [1900, 3730, 3730]
    elif cols == 4:
        widths = [1800, 3260, 1800, 2500]
    else:
        widths = [int(9360 / cols)] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)

    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.15)
                if row_idx == 0 or cell_idx == 0 and cols == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=9.5 if cols >= 4 else 10, color=RGBColor(0, 0, 0), bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, TABLE_FILL)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("项目策划书")
    set_run_font(r, size=16, color=GRAY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("“智途云枢”城市交通智能决策平台")
    set_run_font(r, size=25, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("面向福建宁德城市交通治理的多源数据融合与智能决策服务")
    set_run_font(r, size=12.5, color=GRAY)

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "项目类型"
    table.cell(0, 1).text = "智慧交通类软件项目"
    table.cell(1, 0).text = "技术架构"
    table.cell(1, 1).text = "FastAPI 后端服务 + 前后端分离接口"
    table.cell(2, 0).text = "示范城市"
    table.cell(2, 1).text = "福建宁德"
    table.cell(3, 0).text = "文档日期"
    table.cell(3, 1).text = "2026 年 7 月"
    style_table(table)

    doc.add_page_break()


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("智途云枢项目策划书")
        set_run_font(run, size=9, color=GRAY)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [part.strip() for part in lines[idx].strip().strip("|").split("|")]
        if not all(set(part) <= {"-", ":", " "} for part in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def add_markdown_body(doc, text):
    lines = text.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if idx == 0 and stripped.startswith("# "):
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for r_i, row in enumerate(rows):
                    for c_i, value in enumerate(row):
                        table.cell(r_i, c_i).text = value
                style_table(table)
                doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=4, line=1.208)
            r = p.add_run(stripped[2:])
            set_run_font(r, size=11)
        elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:5]:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, after=4, line=1.208)
            r = p.add_run(stripped.split(". ", 1)[1])
            set_run_font(r, size=11)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, after=8, line=1.333)
            r = p.add_run(stripped.replace("`", ""))
            set_run_font(r, size=11)
        idx += 1


def build():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))
    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
